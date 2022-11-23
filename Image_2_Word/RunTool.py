# ---************************************************---
# @coding: utf-8
# @Time : 2022/11/23 0023 14:53
# @Author : Matrix
# @File : RunTool.py
# @Software: PyCharm
# ---************************************************---
import argparse
import os
from docx import Document
from docx.shared import Inches
from PIL import Image

def Get_Word(args):
    m_images = Get_Image_Path(args.images_path)  # 保存在本地的图片
    m_doc=Document()
    try:
        for i in range(len(m_images)):
            if len(args.image_name) == 0:
                m_doc.add_picture(m_images[i], width=Inches(args.image_width))  # 添加图，设置宽度
                pass
            else:
                str_name = args.image_name + "_" + str(i + 1)
                m_doc.add_paragraph(str_name)  # 添加文字
                m_doc.add_picture(m_images[i], width=Inches(args.image_width))  # 添加图，设置宽度
                pass
            pass
        # m_doc.add_picture(m_images, width=Inches(args.image_width))  # 添加图，设置宽度
        pass
    except Exception:
        for i in range(len(m_images)):
            if len(args.image_name) == 0:
                jpg_ima = Image.open(m_images[i])  # 打开图片
                jpg_ima.save(m_images[i])  # 保存新的图片
                m_doc.add_picture(m_images[i], width=Inches(args.images_width))  # 添加图，设置宽度
                pass
            else:
                str_name = args.image_name + "_" + str(i + 1)
                m_doc.add_paragraph(str_name)  # 添加文字
                jpg_ima = Image.open(m_images[i])  # 打开图片
                jpg_ima.save(m_images[i])  # 保存新的图片
                m_doc.add_picture(m_images[i], width=Inches(args.images_width))  # 添加图，设置宽度
                pass
            pass
        # jpg_ima = Image, open(m_images)  # 打开图片
        # jpg_ima.save(m_images)  # 保存新的图片
        # m_doc.add_picture(m_images, width=Inches(args.image_width))  # 添加图，设置宽度
        pass
    m_doc.save(args.word_path)  # 保存路径
    pass
def Get_Image_Path(path):
    """
    #获取图片路径
    :param path:
    :return:
    """
    image_list=[]
    image_realpath=os.path.abspath(path)
    image_names=os.listdir(image_realpath)
    for image_name in image_names:
        if image_name.endswith('jpg') or image_name.endswith('png'):
            image_path=os.path.join(image_realpath,image_name)
            image_list.append(image_path)
            pass
        pass
    return image_list
    pass
def Get_Titel_Lengeth(image_name):
    if len(image_name)==0:
        return 0
        pass
    else:
        return 1
    pass
def main(args):
    Get_Word(args)
    pass
if __name__=='__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--images_path', type=str, default='./data',required=True, help='图片文件夹地址')
    parser.add_argument('--image_name', type=str, default='',required=False, help='图片名称')
    parser.add_argument('--images_width', type=int, default=4,required=False, help='图片高度')
    parser.add_argument('--word_path', type=str, default='./test.docx',required=False, help='word文件地址')
    args = parser.parse_args()
    main(args)
    pass